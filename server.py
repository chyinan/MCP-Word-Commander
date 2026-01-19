import logging
import os
import json
from typing import List, Optional, Dict, Any, Union
from mcp.server.fastmcp import FastMCP, Image as MCPImage
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.table import Table

# 初始化 MCP Server
mcp = FastMCP("word-commander")

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def _get_abs_path(path: str) -> str:
    """获取绝对路径，支持相对路径"""
    if os.path.isabs(path):
        return path
    return os.path.abspath(os.path.join(os.getcwd(), path))

@mcp.tool()
def create_new_document(file_path: str) -> str:
    """
    创建一个新的空白 Word 文档。
    
    Args:
        file_path: 保存文档的路径 (例如: "output.docx")
    """
    try:
        doc = Document()
        abs_path = _get_abs_path(file_path)
        doc.save(abs_path)
        return f"Successfully created new document at {abs_path}"
    except Exception as e:
        return f"Error creating document: {str(e)}"

@mcp.tool()
def get_document_info(file_path: str) -> str:
    """
    获取 Word 文档的基本信息概览，包括段落数、表格数、各段落的简要信息。
    用于在读取完整内容前快速了解文档结构。
    
    Args:
        file_path: 文档路径
        
    Returns:
        JSON 格式的文档概览信息。
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        # 统计段落信息
        paragraphs_summary = []
        non_empty_count = 0
        for i, para in enumerate(doc.paragraphs):
            text_preview = para.text.strip()[:50] + "..." if len(para.text.strip()) > 50 else para.text.strip()
            if text_preview:
                non_empty_count += 1
                paragraphs_summary.append({
                    "index": i,
                    "preview": text_preview,
                    "style": para.style.name
                })
        
        info = {
            "file_path": abs_path,
            "total_paragraphs": len(doc.paragraphs),
            "non_empty_paragraphs": non_empty_count,
            "total_tables": len(doc.tables),
            "paragraphs_summary": paragraphs_summary
        }
        
        return json.dumps(info, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def read_document_structure(
    file_path: str, 
    start_index: int = 0, 
    limit: int = 50,
    include_empty: bool = False
) -> str:
    """
    分段读取 Word 文档的内容和样式信息。支持分页读取大文档。
    
    Args:
        file_path: 文档路径
        start_index: 从第几个段落开始读取 (0-based)
        limit: 最多读取多少个段落 (默认50)
        include_empty: 是否包含空段落 (默认False)
        
    Returns:
        JSON 格式的字符串，包含段落文本和对应的样式信息。
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        content = []
        total = len(doc.paragraphs)
        count = 0
        
        for i, para in enumerate(doc.paragraphs):
            if i < start_index:
                continue
                
            # 跳过空段落（除非指定包含）
            if not include_empty and not para.text.strip():
                continue
            
            if count >= limit:
                break
                
            style_info = {
                "index": i,
                "text": para.text,
                "style_name": para.style.name,
                "alignment": str(para.alignment) if para.alignment else "LEFT",
                "runs": []
            }
            
            # 深入分析 Run (具体的文字片段) 以获取具体字体信息
            for run in para.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None
                }
                # 尝试获取中文字体设置 (complex script)
                if run._element.rPr is not None:
                    rFonts = run._element.rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        if '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia' in rFonts.attrib:
                            run_info["east_asia_font"] = rFonts.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia']
                
                style_info["runs"].append(run_info)
            
            content.append(style_info)
            count += 1
        
        result = {
            "total_paragraphs": total,
            "returned_count": len(content),
            "start_index": start_index,
            "has_more": (start_index + count) < total,
            "paragraphs": content
        }
            
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@mcp.tool()
def read_tables(file_path: str, table_index: Optional[int] = None) -> str:
    """
    读取 Word 文档中的表格内容。
    
    Args:
        file_path: 文档路径
        table_index: 指定读取第几个表格 (0-based)，不指定则读取所有表格
        
    Returns:
        JSON 格式的表格数据。
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        tables_data = []
        
        for idx, table in enumerate(doc.tables):
            if table_index is not None and idx != table_index:
                continue
                
            table_content = {
                "table_index": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_content["data"].append(row_data)
            
            tables_data.append(table_content)
        
        return json.dumps({
            "total_tables": len(doc.tables),
            "tables": tables_data
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


def _apply_paragraph_format(paragraph, run, font_name: str, font_size: float, 
                            is_bold: bool, alignment: str, indent_first_line: float,
                            line_spacing: Optional[float] = None):
    """内部函数：应用段落格式"""
    align_map = {
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    paragraph.alignment = align_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.LEFT)
    
    if indent_first_line > 0:
        paragraph.paragraph_format.first_line_indent = Pt(font_size * indent_first_line)
    
    if line_spacing:
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    
    run.bold = is_bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


@mcp.tool()
def add_formatted_paragraph(
    file_path: str, 
    text: str, 
    font_name: str = "SimSun", 
    font_size: float = 12.0, 
    is_bold: bool = False,
    alignment: str = "LEFT",
    indent_first_line: float = 0.0,
    line_spacing: Optional[float] = None
) -> str:
    """
    向文档追加带有特定样式的段落。支持设置中文字体（如宋体/黑体）。
    
    Args:
        file_path: 文档路径
        text: 段落内容
        font_name: 字体名称 (默认 "SimSun" 即宋体, 可选 "Microsoft YaHei", "Times New Roman" 等)
        font_size: 字号 (pt)
        is_bold: 是否加粗
        alignment: 对齐方式 ("LEFT", "CENTER", "RIGHT", "JUSTIFY")
        indent_first_line: 首行缩进字符数 (例如 2.0 代表缩进两个字符)
        line_spacing: 行距 (pt)，例如 20.0 代表固定值20磅
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        _apply_paragraph_format(paragraph, run, font_name, font_size, 
                               is_bold, alignment, indent_first_line, line_spacing)
        
        doc.save(abs_path)
        return json.dumps({"success": True, "message": "Successfully added formatted paragraph."}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def replace_paragraph(
    file_path: str,
    paragraph_index: int,
    new_text: str,
    font_name: str = "SimSun",
    font_size: float = 12.0,
    is_bold: bool = False,
    alignment: str = "LEFT",
    indent_first_line: float = 0.0,
    line_spacing: Optional[float] = None
) -> str:
    """
    替换指定段落的内容，并应用新的格式。
    
    Args:
        file_path: 文档路径
        paragraph_index: 要替换的段落索引 (0-based)
        new_text: 新的段落内容
        font_name: 字体名称
        font_size: 字号 (pt)
        is_bold: 是否加粗
        alignment: 对齐方式
        indent_first_line: 首行缩进字符数
        line_spacing: 行距 (pt)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                "error": f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs."
            }, ensure_ascii=False)
        
        para = doc.paragraphs[paragraph_index]
        
        # 清除原有内容
        para.clear()
        
        # 添加新内容
        run = para.add_run(new_text)
        _apply_paragraph_format(para, run, font_name, font_size, 
                               is_bold, alignment, indent_first_line, line_spacing)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True, 
            "message": f"Successfully replaced paragraph at index {paragraph_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def insert_paragraph_after(
    file_path: str,
    after_index: int,
    text: str,
    font_name: str = "SimSun",
    font_size: float = 12.0,
    is_bold: bool = False,
    alignment: str = "LEFT",
    indent_first_line: float = 0.0,
    line_spacing: Optional[float] = None
) -> str:
    """
    在指定段落之后插入新段落。
    
    Args:
        file_path: 文档路径
        after_index: 在此段落索引之后插入 (0-based)
        text: 新段落内容
        font_name: 字体名称
        font_size: 字号 (pt)
        is_bold: 是否加粗
        alignment: 对齐方式
        indent_first_line: 首行缩进字符数
        line_spacing: 行距 (pt)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if after_index < 0 or after_index >= len(doc.paragraphs):
            return json.dumps({
                "error": f"Invalid paragraph index: {after_index}. Document has {len(doc.paragraphs)} paragraphs."
            }, ensure_ascii=False)
        
        # 获取目标段落
        target_para = doc.paragraphs[after_index]
        
        # 在目标段落之后插入新段落
        new_para = OxmlElement('w:p')
        target_para._element.addnext(new_para)
        
        # 创建新的段落对象
        from docx.text.paragraph import Paragraph
        new_paragraph = Paragraph(new_para, target_para._parent)
        
        run = new_paragraph.add_run(text)
        _apply_paragraph_format(new_paragraph, run, font_name, font_size,
                               is_bold, alignment, indent_first_line, line_spacing)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully inserted paragraph after index {after_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def search_and_replace(
    file_path: str,
    search_text: str,
    replace_text: str,
    match_case: bool = True
) -> str:
    """
    在文档中搜索并替换文本（保持原有格式）。
    
    Args:
        file_path: 文档路径
        search_text: 要搜索的文本
        replace_text: 替换为的文本
        match_case: 是否区分大小写 (默认True)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        count = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                if match_case:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)
                        count += 1
                else:
                    import re
                    if re.search(re.escape(search_text), run.text, re.IGNORECASE):
                        run.text = re.sub(re.escape(search_text), replace_text, run.text, flags=re.IGNORECASE)
                        count += 1
        
        # 同时处理表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if match_case:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
                                    count += 1
                            else:
                                import re
                                if re.search(re.escape(search_text), run.text, re.IGNORECASE):
                                    run.text = re.sub(re.escape(search_text), replace_text, run.text, flags=re.IGNORECASE)
                                    count += 1
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Replaced {count} occurrence(s) of '{search_text}'."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)

@mcp.tool()
def create_table_with_data(
    file_path: str,
    rows: int,
    cols: int,
    data: List[List[str]],
    header_bold: bool = True
) -> str:
    """
    在文档末尾创建一个表格并填充数据。
    
    Args:
        file_path: 文档路径
        rows: 行数
        cols: 列数
        data: 二维数组，包含要填充的数据 [['Header1', 'Header2'], ['Val1', 'Val2']]
        header_bold: 第一行是否加粗
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return f"File not found: {abs_path}"
            
        doc = Document(abs_path)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid' # 使用默认带边框的样式
        
        # 填充数据
        for r in range(min(rows, len(data))):
            row_data = data[r]
            for c in range(min(cols, len(row_data))):
                cell = table.cell(r, c)
                cell.text = str(row_data[c])
                
                # 如果是表头且需要加粗
                if r == 0 and header_bold:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            
        doc.save(abs_path)
        return "Successfully created table."
    except Exception as e:
        return f"Error creating table: {str(e)}"


# ============================================
# 图片处理功能
# ============================================

def _extract_images_from_document(doc: Document, include_tables: bool = True) -> List[Dict]:
    """
    内部函数：从文档中提取所有图片信息。
    
    Returns:
        包含图片信息的列表，每个元素包含：
        - index: 图片索引
        - paragraph_index: 所在段落索引
        - location: 位置描述 ("paragraph" 或 "table")
        - image_blob: 图片二进制数据
        - content_type: MIME 类型
        - filename: 原始文件名
        - width_emu: 宽度 (EMU)
        - height_emu: 高度 (EMU)
        - drawing_element: Drawing 的 XML 元素引用
    """
    images = []
    image_idx = 0
    
    # 从段落中提取图片
    for para_idx, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            # 检查 run 中的 drawing 元素
            drawings = run._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for drawing in drawings:
                # 查找 blip 元素获取图片引用
                blips = drawing.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                    if embed_attr in blip.attrib:
                        rId = blip.attrib[embed_attr]
                        try:
                            image_part = doc.part.related_parts.get(rId)
                            if image_part and hasattr(image_part, 'blob'):
                                # 获取尺寸信息
                                extent = drawing.find('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                width_emu = int(extent.get('cx', 0)) if extent is not None else 0
                                height_emu = int(extent.get('cy', 0)) if extent is not None else 0
                                
                                images.append({
                                    'index': image_idx,
                                    'paragraph_index': para_idx,
                                    'location': 'paragraph',
                                    'table_info': None,
                                    'image_blob': image_part.blob,
                                    'content_type': image_part.content_type,
                                    'filename': os.path.basename(image_part.partname),
                                    'width_emu': width_emu,
                                    'height_emu': height_emu,
                                    'width_inches': width_emu / 914400 if width_emu else 0,
                                    'height_inches': height_emu / 914400 if height_emu else 0,
                                    'drawing_element': drawing,
                                    'run_element': run._element,
                                    'rId': rId
                                })
                                image_idx += 1
                        except Exception as e:
                            logger.warning(f"Failed to extract image: {e}")
    
    # 从表格中提取图片
    if include_tables:
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            drawings = run._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            for drawing in drawings:
                                blips = drawing.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                                for blip in blips:
                                    embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                    if embed_attr in blip.attrib:
                                        rId = blip.attrib[embed_attr]
                                        try:
                                            image_part = doc.part.related_parts.get(rId)
                                            if image_part and hasattr(image_part, 'blob'):
                                                extent = drawing.find('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                                width_emu = int(extent.get('cx', 0)) if extent is not None else 0
                                                height_emu = int(extent.get('cy', 0)) if extent is not None else 0
                                                
                                                images.append({
                                                    'index': image_idx,
                                                    'paragraph_index': None,
                                                    'location': 'table',
                                                    'table_info': {
                                                        'table_index': table_idx,
                                                        'row': row_idx,
                                                        'col': cell_idx
                                                    },
                                                    'image_blob': image_part.blob,
                                                    'content_type': image_part.content_type,
                                                    'filename': os.path.basename(image_part.partname),
                                                    'width_emu': width_emu,
                                                    'height_emu': height_emu,
                                                    'width_inches': width_emu / 914400 if width_emu else 0,
                                                    'height_inches': height_emu / 914400 if height_emu else 0,
                                                    'drawing_element': drawing,
                                                    'run_element': run._element,
                                                    'rId': rId
                                                })
                                                image_idx += 1
                                        except Exception as e:
                                            logger.warning(f"Failed to extract image from table: {e}")
    
    return images


@mcp.tool()
def get_images_info(file_path: str, include_tables: bool = True) -> str:
    """
    获取 Word 文档中所有图片的元信息，不返回图片内容。
    用于在读取图片前了解文档中图片的结构和位置。
    
    Args:
        file_path: 文档路径
        include_tables: 是否包含表格中的图片 (默认True)
        
    Returns:
        JSON 格式的图片元信息列表。
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        images = _extract_images_from_document(doc, include_tables)
        
        # 构建不包含二进制数据的信息
        images_info = []
        for img in images:
            info = {
                'index': img['index'],
                'location': img['location'],
                'paragraph_index': img['paragraph_index'],
                'table_info': img['table_info'],
                'filename': img['filename'],
                'content_type': img['content_type'],
                'width_inches': round(img['width_inches'], 2),
                'height_inches': round(img['height_inches'], 2),
                'size_bytes': len(img['image_blob'])
            }
            images_info.append(info)
        
        return json.dumps({
            "total_images": len(images_info),
            "images": images_info
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def read_images(
    file_path: str, 
    image_index: Optional[int] = None,
    include_tables: bool = True
) -> Union[List[MCPImage], str]:
    """
    读取 Word 文档中的图片，直接返回给 AI 查看。
    
    Args:
        file_path: 文档路径
        image_index: 指定读取第几张图片 (0-based)，不指定则读取所有图片
        include_tables: 是否包含表格中的图片 (默认True)
        
    Returns:
        Image 对象列表，AI 可以直接"看到"这些图片。
        如果发生错误，返回错误信息字符串。
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return f"Error: File not found: {abs_path}"
            
        doc = Document(abs_path)
        extracted_images = _extract_images_from_document(doc, include_tables)
        
        if len(extracted_images) == 0:
            return "No images found in the document."
        
        result_images = []
        
        for img in extracted_images:
            if image_index is not None and img['index'] != image_index:
                continue
            
            # 从 content_type 获取格式
            content_type = img['content_type']
            fmt = content_type.split('/')[-1] if '/' in content_type else 'png'
            # 处理特殊情况
            if fmt == 'jpeg':
                fmt = 'jpeg'
            elif fmt not in ['png', 'gif', 'webp', 'jpeg', 'jpg']:
                fmt = 'png'
            
            result_images.append(MCPImage(data=img['image_blob'], format=fmt))
        
        if image_index is not None and len(result_images) == 0:
            return f"Error: Image index {image_index} not found. Document has {len(extracted_images)} images (0-{len(extracted_images)-1})."
        
        return result_images
    except Exception as e:
        return f"Error reading images: {str(e)}"


@mcp.tool()
def add_image(
    file_path: str,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
    alignment: str = "CENTER"
) -> str:
    """
    在文档末尾添加图片。
    
    Args:
        file_path: 文档路径
        image_path: 图片文件路径 (支持 PNG, JPG, GIF, BMP 等格式)
        width_inches: 图片宽度 (英寸)，不指定则使用原始宽度
        height_inches: 图片高度 (英寸)，不指定则按宽度等比缩放
        alignment: 对齐方式 ("LEFT", "CENTER", "RIGHT")
    """
    try:
        abs_path = _get_abs_path(file_path)
        abs_image_path = _get_abs_path(image_path)
        
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"Document not found: {abs_path}"}, ensure_ascii=False)
        if not os.path.exists(abs_image_path):
            return json.dumps({"error": f"Image not found: {abs_image_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        # 添加新段落
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # 设置对齐方式
        align_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT
        }
        paragraph.alignment = align_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)
        
        # 添加图片
        if width_inches:
            if height_inches:
                run.add_picture(abs_image_path, width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(abs_image_path, width=Inches(width_inches))
        elif height_inches:
            run.add_picture(abs_image_path, height=Inches(height_inches))
        else:
            run.add_picture(abs_image_path)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully added image to document."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def insert_image_after_paragraph(
    file_path: str,
    after_index: int,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
    alignment: str = "CENTER"
) -> str:
    """
    在指定段落之后插入图片。
    
    Args:
        file_path: 文档路径
        after_index: 在此段落索引之后插入 (0-based)
        image_path: 图片文件路径
        width_inches: 图片宽度 (英寸)
        height_inches: 图片高度 (英寸)
        alignment: 对齐方式 ("LEFT", "CENTER", "RIGHT")
    """
    try:
        abs_path = _get_abs_path(file_path)
        abs_image_path = _get_abs_path(image_path)
        
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"Document not found: {abs_path}"}, ensure_ascii=False)
        if not os.path.exists(abs_image_path):
            return json.dumps({"error": f"Image not found: {abs_image_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if after_index < 0 or after_index >= len(doc.paragraphs):
            return json.dumps({
                "error": f"Invalid paragraph index: {after_index}. Document has {len(doc.paragraphs)} paragraphs."
            }, ensure_ascii=False)
        
        # 获取目标段落
        target_para = doc.paragraphs[after_index]
        
        # 创建新段落元素
        new_para = OxmlElement('w:p')
        target_para._element.addnext(new_para)
        
        # 创建新的段落对象
        from docx.text.paragraph import Paragraph
        new_paragraph = Paragraph(new_para, target_para._parent)
        
        # 设置对齐方式
        align_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT
        }
        new_paragraph.alignment = align_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)
        
        # 添加图片
        run = new_paragraph.add_run()
        if width_inches:
            if height_inches:
                run.add_picture(abs_image_path, width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(abs_image_path, width=Inches(width_inches))
        elif height_inches:
            run.add_picture(abs_image_path, height=Inches(height_inches))
        else:
            run.add_picture(abs_image_path)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully inserted image after paragraph {after_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def delete_image(file_path: str, image_index: int) -> str:
    """
    删除文档中指定索引的图片。
    使用 get_images_info 工具可以查看所有图片的索引。
    
    Args:
        file_path: 文档路径
        image_index: 要删除的图片索引 (0-based)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        images = _extract_images_from_document(doc, include_tables=True)
        
        if image_index < 0 or image_index >= len(images):
            return json.dumps({
                "error": f"Invalid image index: {image_index}. Document has {len(images)} images (0-{len(images)-1})."
            }, ensure_ascii=False)
        
        target_image = images[image_index]
        drawing_element = target_image['drawing_element']
        
        # 删除 drawing 元素
        parent = drawing_element.getparent()
        if parent is not None:
            parent.remove(drawing_element)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully deleted image at index {image_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def replace_image(
    file_path: str,
    image_index: int,
    new_image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None
) -> str:
    """
    替换文档中指定索引的图片。
    使用 get_images_info 工具可以查看所有图片的索引。
    
    Args:
        file_path: 文档路径
        image_index: 要替换的图片索引 (0-based)
        new_image_path: 新图片文件路径
        width_inches: 新图片宽度 (英寸)，不指定则使用原图片尺寸
        height_inches: 新图片高度 (英寸)
    """
    try:
        abs_path = _get_abs_path(file_path)
        abs_new_image_path = _get_abs_path(new_image_path)
        
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"Document not found: {abs_path}"}, ensure_ascii=False)
        if not os.path.exists(abs_new_image_path):
            return json.dumps({"error": f"New image not found: {abs_new_image_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        images = _extract_images_from_document(doc, include_tables=True)
        
        if image_index < 0 or image_index >= len(images):
            return json.dumps({
                "error": f"Invalid image index: {image_index}. Document has {len(images)} images (0-{len(images)-1})."
            }, ensure_ascii=False)
        
        target_image = images[image_index]
        drawing_element = target_image['drawing_element']
        run_element = target_image['run_element']
        
        # 如果没有指定尺寸，使用原图片尺寸
        if width_inches is None:
            width_inches = target_image['width_inches'] if target_image['width_inches'] > 0 else None
        if height_inches is None:
            height_inches = target_image['height_inches'] if target_image['height_inches'] > 0 else None
        
        # 找到包含该 drawing 的 run
        from docx.text.run import Run
        run = Run(run_element, None)
        
        # 删除旧的 drawing
        parent = drawing_element.getparent()
        if parent is not None:
            parent.remove(drawing_element)
        
        # 在同一个 run 中添加新图片
        if width_inches:
            if height_inches:
                run.add_picture(abs_new_image_path, width=Inches(width_inches), height=Inches(height_inches))
            else:
                run.add_picture(abs_new_image_path, width=Inches(width_inches))
        elif height_inches:
            run.add_picture(abs_new_image_path, height=Inches(height_inches))
        else:
            run.add_picture(abs_new_image_path)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully replaced image at index {image_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


# ============================================
# 表格增强功能
# ============================================

@mcp.tool()
def insert_table_after_paragraph(
    file_path: str,
    after_index: int,
    rows: int,
    cols: int,
    data: List[List[str]],
    header_bold: bool = True,
    style: str = "Table Grid"
) -> str:
    """
    在指定段落之后插入表格。
    
    Args:
        file_path: 文档路径
        after_index: 在此段落索引之后插入表格 (0-based)
        rows: 表格行数
        cols: 表格列数
        data: 二维数组，包含要填充的数据 [['Header1', 'Header2'], ['Val1', 'Val2']]
        header_bold: 第一行是否加粗 (默认True)
        style: 表格样式 (默认 "Table Grid")
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if after_index < 0 or after_index >= len(doc.paragraphs):
            return json.dumps({
                "error": f"Invalid paragraph index: {after_index}. Document has {len(doc.paragraphs)} paragraphs."
            }, ensure_ascii=False)
        
        # 获取目标段落
        target_para = doc.paragraphs[after_index]
        
        # 创建表格
        table = doc.add_table(rows=rows, cols=cols)
        try:
            table.style = style
        except:
            table.style = 'Table Grid'  # 如果指定样式不存在，使用默认样式
        
        # 填充数据
        for r in range(min(rows, len(data))):
            row_data = data[r]
            for c in range(min(cols, len(row_data))):
                cell = table.cell(r, c)
                cell.text = str(row_data[c])
                
                # 如果是表头且需要加粗
                if r == 0 and header_bold:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # 将表格移动到目标段落之后
        target_para._element.addnext(table._tbl)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully inserted {rows}x{cols} table after paragraph {after_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def update_table_cell(
    file_path: str,
    table_index: int,
    row: int,
    col: int,
    new_text: str,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    is_bold: Optional[bool] = None
) -> str:
    """
    修改表格中指定单元格的内容。
    
    Args:
        file_path: 文档路径
        table_index: 表格索引 (0-based)
        row: 行索引 (0-based)
        col: 列索引 (0-based)
        new_text: 新的单元格内容
        font_name: 字体名称 (可选)
        font_size: 字号 pt (可选)
        is_bold: 是否加粗 (可选)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if table_index < 0 or table_index >= len(doc.tables):
            return json.dumps({
                "error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."
            }, ensure_ascii=False)
        
        table = doc.tables[table_index]
        
        if row < 0 or row >= len(table.rows):
            return json.dumps({
                "error": f"Invalid row index: {row}. Table has {len(table.rows)} rows."
            }, ensure_ascii=False)
        
        if col < 0 or col >= len(table.columns):
            return json.dumps({
                "error": f"Invalid column index: {col}. Table has {len(table.columns)} columns."
            }, ensure_ascii=False)
        
        cell = table.cell(row, col)
        
        # 清除原有内容并设置新文本
        cell.text = new_text
        
        # 应用格式
        if font_name or font_size is not None or is_bold is not None:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if font_name:
                        run.font.name = font_name
                        # 设置中文字体
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), font_name)
                    if font_size is not None:
                        run.font.size = Pt(font_size)
                    if is_bold is not None:
                        run.bold = is_bold
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully updated cell ({row}, {col}) in table {table_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def delete_table(file_path: str, table_index: int) -> str:
    """
    删除文档中指定索引的表格。
    
    Args:
        file_path: 文档路径
        table_index: 要删除的表格索引 (0-based)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if table_index < 0 or table_index >= len(doc.tables):
            return json.dumps({
                "error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."
            }, ensure_ascii=False)
        
        table = doc.tables[table_index]
        tbl = table._tbl
        tbl.getparent().remove(tbl)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully deleted table at index {table_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def add_table_row(
    file_path: str,
    table_index: int,
    row_data: List[str],
    position: Optional[int] = None
) -> str:
    """
    向表格中添加新行。
    
    Args:
        file_path: 文档路径
        table_index: 表格索引 (0-based)
        row_data: 新行的数据列表
        position: 插入位置 (0-based)，不指定则添加到末尾
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if table_index < 0 or table_index >= len(doc.tables):
            return json.dumps({
                "error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."
            }, ensure_ascii=False)
        
        table = doc.tables[table_index]
        
        # 添加新行
        if position is not None:
            if position < 0 or position > len(table.rows):
                return json.dumps({
                    "error": f"Invalid position: {position}. Table has {len(table.rows)} rows."
                }, ensure_ascii=False)
            # 在指定位置插入行
            new_row = table.add_row()
            # 移动到正确位置
            tbl = table._tbl
            tr = new_row._tr
            tbl.remove(tr)
            tbl.insert(position + 1, tr)  # +1 因为第一个元素是 tblPr
        else:
            new_row = table.add_row()
        
        # 填充数据
        for i, cell_text in enumerate(row_data):
            if i < len(new_row.cells):
                new_row.cells[i].text = str(cell_text)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully added row to table {table_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


@mcp.tool()
def delete_table_row(file_path: str, table_index: int, row_index: int) -> str:
    """
    删除表格中的指定行。
    
    Args:
        file_path: 文档路径
        table_index: 表格索引 (0-based)
        row_index: 要删除的行索引 (0-based)
    """
    try:
        abs_path = _get_abs_path(file_path)
        if not os.path.exists(abs_path):
            return json.dumps({"error": f"File not found: {abs_path}"}, ensure_ascii=False)
            
        doc = Document(abs_path)
        
        if table_index < 0 or table_index >= len(doc.tables):
            return json.dumps({
                "error": f"Invalid table index: {table_index}. Document has {len(doc.tables)} tables."
            }, ensure_ascii=False)
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            return json.dumps({
                "error": f"Invalid row index: {row_index}. Table has {len(table.rows)} rows."
            }, ensure_ascii=False)
        
        # 删除行
        tr = table.rows[row_index]._tr
        table._tbl.remove(tr)
        
        doc.save(abs_path)
        return json.dumps({
            "success": True,
            "message": f"Successfully deleted row {row_index} from table {table_index}."
        }, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


def main():
    """Entry point for the MCP Word Commander server."""
    mcp.run()


if __name__ == "__main__":
    main()

